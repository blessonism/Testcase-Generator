"""
@filename:WordLocat.py
@author:JunLeon
@time:2023-10-23
"""
import os
import re
# 定位word中的表格框值
from docx import Document

from format import format_table_content, format_table_text
from split import extract_and_map_text

from docx import Document


# 打开Word文档
doc = Document('Test-case.docx')

def replace_table_value(doc, old_value, new_value):
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if old_value in paragraph.text:
                        paragraph.text = paragraph.text.replace(old_value, new_value)



# input_text = """
# id_value: 6
# Module_Name_value: Div module;
# Test_Title_value: User enter invalid number and we get an error;
# Description_value: User enters invalid numbers for division and expects an error;
# pre-conditions_value: User is attempting to divide two invalid numbers;
# Test Step1: User enters one invalid number for division;
# Test Data1: Abc
# Expected Result1: Get error tips;
# Actual Result1: We get error tips;
# Test Step2: User enters two invalid numbers for division;
# Test Data2: xyz
# Expected Result2: Get error tips;
# Actual Result2: We get error tips;
# Test Step3: User enters "/" operator for division;
# Test Data3: /
# Expected Result3: Get warn tip;
# Actual Result3: We get warn tip;
# Test Step4: User presses the "=" operator to perform division;
# Test Data4: =
# Expected Result4: Get warn tip;
# Actual Result4: We get warn tip;
# Post-conditions_value: If the numbers are invalid, we can't divide them, and our test case is pass.
#
# """

input_text = ""
try:
    while True:
        line = input()
        if not line:
            break
        input_text += line + "\n"
except EOFError:
    pass

# 查找并提取id_value后的值
match = re.search(r'id_value:\s*(\d+)', input_text)

value_id = ""

if match:
    value_id = match.group(1)

# 查找并提取Module_Name_value后的值
match = re.search(r'Module_Name_value:\s*([^;]+)', input_text)

if match:
    value_name = match.group(1).strip()


outDocName = "Test-case_" + value_id + "_" + value_name

result = extract_and_map_text(input_text)


for key, value in result.items():
    # 调用方法来替换文本
    replace_table_value(doc, key, value)

    # print(f"map[{key}] = \"{value}\"")

# 清空预设模板中多余的步骤
def clear_rows_with_text(table, text_to_clear):
    for row in table.rows:
        cell_texts = [cell.text.strip() for cell in row.cells]
        if any(keyword in cell_texts for keyword in text_to_clear):
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    paragraph.clear()


# 要清空的文本列表
text_to_clear = ["Test Step2", "Test Step3", "Test Step4", "Test Step5"]

for table in doc.tables:
    clear_rows_with_text(table, text_to_clear)

# 居中表格数据
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.alignment = 1  # 1表示居中对齐



# 调用方法来格式化第一个表格中的文本
format_table_text(doc.tables)

# 调用方法来格式化第二个表格中的文本
format_table_content(doc.tables)


output_folder = 'TestCase'

# Ensure the folder exists, or create it if it doesn't
if not os.path.exists(output_folder):
    os.makedirs(output_folder)

# Construct the output file path
output_file = os.path.join(output_folder, outDocName + '.docx')

# Save the modified document to the TestCase folder
doc.save(output_file)
print("finished")