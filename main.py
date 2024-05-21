"""
@filename: main.py
@author: Edom
@time:2024-5-21
"""
import os
import re
from docx import Document
from format import format_table_content, format_table_text
from split import extract_and_map_text


def replace_table_value(doc, old_value, new_value):
    """替换Word文档表格中的文本"""
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if old_value in paragraph.text:
                        paragraph.text = paragraph.text.replace(old_value, new_value)


def clear_rows_with_text(table, text_to_clear):
    """清除包含特定文本的行"""
    rows_to_clear = []
    for row in table.rows:
        cell_texts = [cell.text.strip() for cell in row.cells]
        if any(keyword in cell_texts for keyword in text_to_clear):
            rows_to_clear.append(row)

    for row in rows_to_clear:
        table._element.remove(row._element)


def remove_extra_steps(doc, num_steps):
    """根据用户输入的测试步骤数，删除多余的行"""
    step_keywords = [f"Test Step{i}" for i in range(num_steps + 1, 9)]
    for table in doc.tables:
        clear_rows_with_text(table, step_keywords)


def center_align_table_text(doc):
    """将所有表格中的文本居中对齐"""
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.alignment = 1  # 1表示居中对齐


def get_input_text():
    """读取用户的多行输入"""
    input_text = ""
    try:
        while True:
            line = input()
            if not line:
                break
            input_text += line + "\n"
    except EOFError:
        pass
    return input_text


def extract_value(pattern, text):
    """使用正则表达式提取值"""
    match = re.search(pattern, text)
    if match:
        return match.group(1).strip()
    return ""


def validate_input(input_text):
    """验证用户输入的规范性"""
    errors = []

    patterns = {
        "id_value": r'id_value:\s*\d+',
        "Module_Name_value": r'Module_Name_value:\s*[^;]+;',
        "Author_value": r'Author_value:\s*[^;]+;',
        "Date_value": r'Date_value:\s*\d{4}\.\d{2}\.\d{2};',
        "Test_Title_value": r'Test_Title_value:\s*[^;]+;',
        "Description_value": r'Description_value:\s*[^;]+;',
        "pre-conditions_value": r'pre-conditions_value:\s*[^;]+;',
        "Test_Step": r'Test Step\d+:\s*[^;]+;',
        "Test_Data": r'Test Data\d+:\s*[^;]+',
        "Expected_Result": r'Expected Result\d+:\s*[^;]+;',
        "Actual_Result": r'Actual Result\d+:\s*[^;]+;',
        "Post_conditions_value": r'Post-conditions_value:\s*[^;.]+[;.]'
    }

    for key, pattern in patterns.items():
        if key.startswith("Test_Step") or key.startswith("Test_Data") or key.startswith(
                "Expected_Result") or key.startswith("Actual_Result"):
            matches = re.findall(pattern, input_text)
            if not matches:
                errors.append(f"缺少或不符合规范的 {key} 条目。")
        else:
            if not re.search(pattern, input_text):
                errors.append(f"缺少或不符合规范的 {key} 条目。")

    return errors


def main():
    # 获取用户输入文本
    input_text = get_input_text()

    # 验证用户输入
    errors = validate_input(input_text)
    if errors:
        for error in errors:
            print(error)
        return

    # 打开Word文档
    try:
        doc = Document('Testcase_Template.docx')
    except Exception as e:
        print(f"无法打开文档: {e}")
        return

    # 提取id_value和Module_Name_value
    value_id = extract_value(r'id_value:\s*(\d+)', input_text)
    value_name = extract_value(r'Module_Name_value:\s*([^;]+?);', input_text)

    if not value_id or not value_name:
        print("缺少id_value或Module_Name_value")
        return

    outDocName = f"Test-case_{value_id}_{value_name}"

    # 统计用户输入的测试步骤数
    num_steps = len(re.findall(r'Test Step\d+:', input_text))

    print(f'输入数据中包含 Test Step 数为：{num_steps}')
    # 根据输入步骤数删除多余行
    remove_extra_steps(doc, num_steps)

    # 提取并映射文本
    result = extract_and_map_text(input_text)

    # 替换文档中的文本
    for key, value in result.items():
        replace_table_value(doc, key, value)

    # 表格文本居中对齐
    center_align_table_text(doc)

    # 格式化表格文本和内容
    format_table_text(doc.tables)
    format_table_content(doc.tables)

    output_folder = 'TestCase'

    # 确保目录存在
    os.makedirs(output_folder, exist_ok=True)

    # 保存修改后的文档
    output_file = os.path.join(output_folder, f"{outDocName}.docx")
    try:
        doc.save(output_file)
        print(f'输出路径为：{output_file}')
        print("finished")
    except Exception as e:
        print(f"无法保存文档: {e}")
        print("请确保文档未被打开。")


if __name__ == "__main__":
    main()
