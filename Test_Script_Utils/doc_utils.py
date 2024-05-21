"""
@filename: main.py
@author: Edom
@time:2024-5-21
"""

from docx.shared import Pt, RGBColor


def replace_table_value(doc, old_value, new_value):
    """替换Word文档表格中的文本"""
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if old_value in paragraph.text:
                        paragraph.text = paragraph.text.replace(old_value, new_value)


def clear_rows_with_text(table, text_to_clear):
    """清除包含特定文本的行"""
    rows_to_clear = []
    for row in table.rows:
        cell_texts = [cell.text.strip() for cell in row.cells]
        if any(keyword in cell_texts for keyword in text_to_clear):
            rows_to_clear.append(row)

    for row in rows_to_clear:
        table._element.remove(row._element)


def remove_extra_steps(doc, num_steps):
    """根据用户输入的测试步骤数，删除多余的行"""
    step_keywords = [f"Test Step{i}" for i in range(num_steps + 1, 9)]
    for table in doc.tables:
        clear_rows_with_text(table, step_keywords)


def center_align_table_text(doc):
    """将所有表格中的文本居中对齐"""
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.alignment = 1  # 1表示居中对齐


def format_table_text(table):
    """格式化表格中的文本"""
    if table:
        # 获取第一个表格
        first_table = table[0]

        for row in first_table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    # 切分文本以 ":" 为分隔符
                    parts = paragraph.text.split(':')

                    if len(parts) == 2:
                        # 创建一个新文本段落
                        paragraph.clear()

                        # 添加加粗文本
                        run_bold = paragraph.add_run(parts[0].strip())
                        run_bold.bold = True

                        # 添加绿色文本
                        run_green = paragraph.add_run(': ' + parts[1].strip())
                        run_green.font.color.rgb = RGBColor(0x00, 0x80, 0x00)  # 设置为绿色


def format_table_content(table):
    """格式化表格内容"""
    if table and len(table) > 1:
        second_table = table[1]
        for row_index, row in enumerate(second_table.rows):
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    # Skip the first row
                    if row_index > 0:
                        for run in paragraph.runs:
                            # Set font name to Times New Roman and size to 12 pt
                            run.font.name = 'Times New Roman'
                            run.font.size = Pt(12)

                    # Set text alignment to center
                    for tab in paragraph.runs:
                        tab.alignment = 1  # 1 corresponds to center alignment
