"""
@filename:format.py
@author:JunLeon
@time:2023-10-23
"""

from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement



def format_table_text(table):
    if table:
        # 获取第一个表格
        first_table = table[0]

        for row in first_table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    # 切分文本以 ":" 为分隔符
                    parts = paragraph.text.split(':')

                    if len(parts) == 2:
                        # 创建一个新文本段落
                        paragraph.clear()

                        # 添加加粗文本
                        run_bold = paragraph.add_run(parts[0].strip())
                        run_bold.bold = True

                        # 添加绿色文本
                        run_green = paragraph.add_run(': ' + parts[1].strip())
                        run_green.font.color.rgb = RGBColor(0x00, 0x80, 0x00)  # 设置为绿色


def format_table_content(table):
    if table and len(table) > 1:
        second_table = table[1]
        for row_index, row in enumerate(second_table.rows):
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    # Skip the first row
                    if row_index > 0:
                        for run in paragraph.runs:
                            # Set font name to Times New Roman and size to 4 pt
                            run.font.name = 'Times New Roman'
                            run.font.size = Pt(12)

                    # Set text alignment to center
                    for tab in paragraph.runs:
                        tab.alignment = 1  # 1 corresponds to center alignment


